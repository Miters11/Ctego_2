from docx import Document
from docx.oxml import OxmlElement
from docx.shared import *
import os

russian_encoding_ascii = {
    'А': 192, 'Б': 193, 'В': 194, 'Г': 195, 'Д': 196, 'Е': 197, 'Ж': 198,
    'З': 199, 'И': 200, 'Й': 201, 'К': 202, 'Л': 203, 'М': 204, 'Н': 205,
    'О': 206, 'П': 207, 'Р': 208, 'С': 209, 'Т': 210, 'У': 211, 'Ф': 212,
    'Х': 213, 'Ц': 214, 'Ч': 215, 'Ш': 216, 'Щ': 217, 'Ъ': 218, 'Ы': 219,
    'Ь': 220, 'Э': 221, 'Ю': 222, 'Я': 223, 'а': 224,
    'б': 225, 'в': 226, 'г': 227, 'д': 228, 'е': 229, 'ж': 230, 'з': 231,
    'и': 232, 'й': 233, 'к': 234, 'л': 235, 'м': 236, 'н': 237, 'о': 238,
    'п': 239, 'р': 240, 'с': 241, 'т': 242, 'у': 243, 'ф': 244, 'х': 245,
    'ц': 246, 'ч': 247, 'ш': 248, 'щ': 249, 'ъ': 250, 'ы': 251, 'ь': 252,
    'э': 253, 'ю': 254, 'я': 255
}

russian_decoding = {v: k for k, v in russian_encoding_ascii.items()}

ArrayBytesCodRule = ["11111", "00000", "11011", "00100", "00010", "01000"]
ArrayBytesCod = ["00011", "11001", "01110", "01001", "00001", "01101", "11010", "10100", "00110",
                 "01011", "01111", "10010", "11100", "01100", "11000", "10110", "10111", "01010",
                 "00101", "10000", "00111", "11110", "10011", "11101", "10101", "10001"]
ArrayLatinUp = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M',
                'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
ArrayRussianUp = ['А', 'Б', 'Ц', 'Д', 'Е', 'Ф', 'Г', 'Х', 'И', 'Й', 'К', 'Л', 'М',
                  'Н', 'О', 'П', 'Я', 'Р', 'С', 'Т', 'У', 'Ж', 'В', 'Ь', 'Ы', 'З']
ArraySpecialUp = ['-', '?', ':', '', '3', 'Э', 'Ш', 'Щ', '8', 'Ю', '(', ')', '.', ',',
                  '9', '0', '1', '4', '\'', '5', '7', '=', '2', '/', '6', '+']

namespace = {
    'ns0': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}
names = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def decode_on_ascii(binary_string):
    text = ''
    if len(binary_string) % 8 != 0:
        binary_string = binary_string + '0' * (len(binary_string) % 8)
    for i in range(0, len(binary_string), 8):
        chisl = binary_string[i:i + 8]
        if int(chisl, 2) in russian_decoding:
            text += russian_decoding[int(chisl, 2)]
        else:
            text += bytes([int(chisl, 2)]).decode('utf-8', errors='replace')
    text = text.replace('\x00', '')
    return text


def mtk2_decode(binary_string):
    text = ''
    if len(binary_string) % 5 != 0:
        binary_string += '0' * (len(binary_string) % 5)
    flag = 0 if binary_string[:5] == '00000' else 1 if binary_string[:5] == '11111' else 2
    for i in range(5, len(binary_string[5:]), 5):
        char_kod = binary_string[i:i + 5]
        if char_kod == ArrayBytesCodRule[1]:
            flag = 0
        elif char_kod == ArrayBytesCodRule[0]:
            flag = 1
        elif char_kod == ArrayBytesCodRule[2]:
            flag = 2
        elif char_kod == ArrayBytesCodRule[3]:
            text += ' '
        elif char_kod == ArrayBytesCodRule[4] or char_kod == ArrayBytesCodRule[5]:
            text += '\n'
        else:
            ind = ArrayBytesCod.index(char_kod)
            if flag == 0:
                text += ArrayRussianUp[ind]
            elif flag == 1:
                text += ArrayLatinUp[ind]
            else:
                text += ArraySpecialUp[ind]
    return text


def inf_o_nab_simv(run):
    r_element = run._element
    color_elem = r_element.find('.//ns0:color', namespaces=namespace)
    color = color_elem.get('{' + names + '}val') if color_elem is not None else '000000'
    shd_elem = r_element.find('.//ns0:shd', namespaces=namespace)
    background_color = shd_elem.get('{' + names + '}fill') if shd_elem is not None else 'FFFFFF'
    sz_elem = r_element.find('.//ns0:sz', namespaces=namespace)
    font_size = int(sz_elem.get('{' + names + '}val')) / 2 if sz_elem is not None else 'не установлен'
    w_elem = r_element.find('.//ns0:w', namespaces=namespace)
    scale = float(w_elem.get('{' + names + '}val')) if w_elem is not None else 100
    spacing_elem = r_element.find('.//ns0:spacing', namespaces=namespace)
    spacing = int(spacing_elem.get('{' + names + '}val')) / 2 if spacing_elem is not None else 0
    rFonts_elem = r_element.find('.//ns0:rFonts', namespaces=namespace)
    font_name = rFonts_elem.get('{' + names + '}ascii') if rFonts_elem is not None else 'не установлен'
    return [hex_decimal(color), hex_decimal(background_color), font_size, scale, spacing, font_name]


def hex_decimal(hex_string):
    if len(hex_string) == 6 and all(c in '0123456789ABCDEFabcdef' for c in hex_string):
        return ' '.join(str(int(hex_string[i:i + 2], 16)) for i in range(0, 6, 2))
    else:
        decimal_numbers = hex_string.split()
        if len(decimal_numbers) == 3 and all(num.isdigit() for num in decimal_numbers):
            return ''.join(f"{int(num):02X}" for num in decimal_numbers)
        else:
            raise ValueError(
                "Входная строка должна быть либо 6-символьной шестнадцатеричной, либо состоять из 3 десятичных чисел.")


def sverka(tek, etal):
    if tek != etal:
        return 1
    else:
        return 0


def lsttoset(lst):
    unique_sublists = {tuple(str(sublist)) for sublist in lst}
    return len(unique_sublists) == 1


def invert_binary_list(binary_list):
    return [1 if x == 0 else 0 for x in binary_list]


def binary_string(binary_string):
    byte_array = bytearray()
    for i in range(0, len(binary_string), 8):
        byte = binary_string[i:i + 8]
        if len(byte) == 8:
            byte_array.append(int(byte, 2))
    byte_array = bytearray(b if b != 0 else 0x20 for b in byte_array)
    return bytes(byte_array)


def decode_bin_to_text(oglav, text):
    global bodo
    result_decoing={}
    if text[-1] != 0:
        text = invert_binary_list(text)
    stroka = oglav + text
    stroka1 = invert_binary_list(oglav) + text
    result = []
    if ''.join(map(str, stroka[:5])) == '00000' or ''.join(map(str, stroka[:5])) == '11011':
        result_decoing['1. Бодо (МТК-2)']=mtk2_decode("".join(map(str, stroka)))
        bodo=stroka
    elif ''.join(map(str, stroka1[:5])) == '00000' or ''.join(map(str, stroka1[:5])) == '11011':
        result_decoing['1. Бодо (МТК-2)']=mtk2_decode("".join(map(str, stroka1)))
        bodo=stroka1
    if stroka[0] == 1:
        result_decoing['2. KOI8-R']=binary_string("".join(map(str, stroka))).decode('KOI8-R')
        result_decoing['3. cp866'] = binary_string("".join(map(str, stroka))).decode('cp866')
        result_decoing['4. cp1251'] = binary_string("".join(map(str, stroka))).decode('cp1251', errors='replace')
        result_decoing['5. ASCII'] = decode_on_ascii("".join(map(str, stroka)))
        result=stroka.copy()
    elif stroka1[0] == 1:
        result_decoing['2. KOI8-R'] = binary_string("".join(map(str, stroka1))).decode('KOI8-R')
        result_decoing['3. cp866'] = binary_string("".join(map(str, stroka1))).decode('cp866')
        result_decoing['4. cp1251'] = binary_string("".join(map(str, stroka1))).decode('cp1251', errors='replace')
        result_decoing['5. ASCII'] = decode_on_ascii("".join(map(str, stroka1)))
        result = stroka1.copy()
    if result_decoing=={}:
        exit('Ошибка декодирования')
    return result_decoing,result


def interval(doc_path):
    global vibor_dek
    documen,documen1 = [],[]
    interv_sp,interv_sp1 = [],[]
    size_sp,size_sp1 = [],[]
    color_sp,color_sp1 = [],[]
    back_color_sp,back_color_sp1 = [],[]
    masht_sp,masht_sp1 = [],[]

    doc = Document(doc_path)
    zagolovok_style = inf_o_nab_simv(doc.paragraphs[0].runs[0])
    text_style = next((inf_o_nab_simv(paragraph.runs[0]) for paragraph in doc.paragraphs[1:] if paragraph.runs), None)
    for run in doc.paragraphs[0].runs:
        inf_o_sim = inf_o_nab_simv(run)
        bit_int = sverka(zagolovok_style[-2], inf_o_sim[-2])
        bit_size = sverka(zagolovok_style[2], inf_o_sim[2])
        bit_color = sverka(zagolovok_style[0], inf_o_sim[0])
        bit_back_color = sverka(zagolovok_style[1], inf_o_sim[1])
        bit_scale = sverka(zagolovok_style[-3], inf_o_sim[-3])
        n = len(list(run.text))
        documen.extend(list(run.text))
        interv_sp.extend([bit_int] * n)
        size_sp.extend([bit_size] * n)
        color_sp.extend([bit_color] * n)
        back_color_sp.extend([bit_back_color] * n)
        masht_sp.extend([bit_scale] * n)

    for para in doc.paragraphs[1:]:
        for run in para.runs:
            inf_o_sim = inf_o_nab_simv(run)
            if inf_o_sim[2]=='не установлен':
                continue
            bit_int = sverka(text_style[-2], inf_o_sim[-2])
            bit_size = sverka(text_style[2], inf_o_sim[2])
            bit_color = sverka(text_style[0], inf_o_sim[0])
            bit_back_color = sverka(text_style[1], inf_o_sim[1])
            bit_scale = sverka(text_style[-3], inf_o_sim[-3])
            n = len(list(run.text))
            documen1.extend(list(run.text))
            interv_sp1.extend([bit_int] * n)
            size_sp1.extend([bit_size] * n)
            color_sp1.extend([bit_color] * n)
            back_color_sp1.extend([bit_back_color] * n)
            masht_sp1.extend([bit_scale] * n)

    params=()
    if 1 in masht_sp1:
        print('Для кодировки изменяли масштаб текста')
        params = (masht_sp, masht_sp1)

    elif 1 in interv_sp1:
        print('Для кодировки изменяли межсимвольный интервал')
        params = (interv_sp, interv_sp1)
    elif 1 in color_sp1:
        print('Для кодировки изменяли цвет текста')
        params = (color_sp, color_sp1)
    elif 1 in back_color_sp1:
        print('Для кодировки изменяли задний фон текста')
        params = (back_color_sp, back_color_sp1)
    elif 1 in size_sp1:
        print('Для кодировки изменяли размер текста')
        params=(size_sp, size_sp1)
    else:
        exit('Ошибка жизни')
    result_decoding,stroka=decode_bin_to_text(params[0],params[1])
    print(stroka)
    vibor_dek=len(result_decoding)
    print('Результат декодирования:')
    for key in result_decoding.keys():
        print(f'{key}: {result_decoding[key]}')
    return stroka


def write_new_doc(path_for_doc,input_doc,binary_string):
    doc=Document(input_doc)
    new_doc=Document()
    ind=0
    for para in doc.paragraphs:
        new_para=new_doc.add_paragraph()
        for run in para.runs:
            #if inf_o_sim[2]=='не установлен':
                #continue
            color, background_color, font_size, scale, spacing, font = inf_o_nab_simv(run)
            run_text = run.text
            run_length = len(run_text)
            new_run=new_para.add_run(run_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            if 1 in binary_string[ind:ind+run_length]:
                new_run.font.underline = True
            else:
                new_run.font.underline = False
            new_run.font.strike = run.font.strike
            new_run.font.color.rgb = run.font.color.rgb
            r_element = new_run._element

            if r_element.rPr is None:
                rPr = OxmlElement('w:rPr')
                r_element.append(rPr)
            else:
                rPr = r_element.rPr

            if spacing:
                spacing_elem = OxmlElement('w:spacing')
                spacing_elem.set('{' + names + '}val', str(int(spacing) * 2))
                rPr.append(spacing_elem)
            if scale:
                scale_elem = OxmlElement('w:w')
                scale_elem.set('{' + names + '}val', str(scale))
                rPr.append(scale_elem)
            if background_color:
                back_elem = OxmlElement('w:shd')
                back_elem.set('{' + names + '}fill', background_color)
                rPr.append(back_elem)
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            ind+=run_length

    try:
        new_doc.save(path_for_doc)
    except PermissionError as e:
        print(f"Ошибка при сохранении документа: {e}")
    os.startfile(path_for_doc)

def get_spacing(doc_path):
    stroka_bit=interval(doc_path)
    while True:
        vibor=int(input('Выберите результат декодирования: '))
        if vibor<=5 and vibor>0:
            break
    print(stroka_bit)
    if vibor_dek==5 and vibor==1:
        if bodo:
            stroka_bit=bodo
        else:
            exit('Выберите другую кодировку для вывода')
    print(stroka_bit)
    write_new_doc('result_decode.docx',doc_path,stroka_bit)



get_spacing('result.docx')
