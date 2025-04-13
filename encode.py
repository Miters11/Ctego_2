from docx import Document
from docx.shared import RGBColor, Pt
import os
from docx.oxml import OxmlElement

russian_encoding_ascii = {
    'А': 192, 'Б': 193, 'В': 194, 'Г': 195, 'Д': 196, 'Е': 197, 'Ж': 198,
    'З': 199, 'И': 200, 'Й': 201, 'К': 202, 'Л': 203, 'М': 204, 'Н': 205,
    'О': 206, 'П': 207, 'Р': 208, 'С': 209, 'Т': 210, 'У': 211, 'Ф': 212,
    'Х': 213, 'Ц': 214, 'Ч': 215, 'Ш': 216, 'Щ': 217, 'Ъ': 218, 'Ы': 219,
    'Ь': 220, 'Э': 221, 'Ю': 222, 'Я': 223, 'а': 224,
    'б': 225, 'в': 226, 'г': 227, 'д': 228, 'е': 229, 'ж': 230, 'з': 231,
    'и': 232, 'й': 233, 'к': 234, 'л': 235, 'м': 236, 'н': 237, 'о': 238,
    'п': 239, 'р': 240, 'с': 241, 'т': 242, 'у': 243, 'ф': 244, 'х': 245,
    'ц': 246, 'ч': 247, 'ш': 248, 'щ': 249, 'ъ': 250, 'ы': 251, 'ь': 252,
    'э': 253, 'ю': 254, 'я': 255
}

ArrayBytesCodRule = ["11111", "00000", "11011", "00100", "00010", "01000"]
ArrayBytesCod = ["00011", "11001", "01110", "01001", "00001", "01101", "11010", "10100", "00110",
                 "01011", "01111", "10010", "11100", "01100", "11000", "10110", "10111", "01010",
                 "00101", "10000", "00111", "11110", "10011", "11101", "10101", "10001"]
ArrayLatinUp = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M',
                'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
ArrayRussianUp = ['А', 'Б', 'Ц', 'Д', 'Е', 'Ф', 'Г', 'Х', 'И', 'Й', 'К', 'Л', 'М',
                  'Н', 'О', 'П', 'Я', 'Р', 'С', 'Т', 'У', 'Ж', 'В', 'Ь', 'Ы', 'З']
ArraySpecialUp = ['-', '?', ':', '', '3', 'Э', 'Ш', 'Щ', '8', 'Ю', '(', ')', '.', ',',
                  '9', '0', '1', '4', '\'', '5', '7', '=', '2', '/', '6', '+']


def mtk2_encode(text):
    binary_string = '00000' if text[0] in ArrayRussianUp else '11111' if text[0] in ArrayLatinUp else '11011'
    flag = 0 if binary_string == '00000' else 1 if binary_string == '11111' else 2
    ind = 0
    while ind < len(text):
        if flag == 0 and text[ind] in ArrayRussianUp:
            binary_string += ArrayBytesCod[ArrayRussianUp.index(text[ind])]
        elif flag == 1 and text[ind] in ArrayLatinUp:
            binary_string += ArrayBytesCod[ArrayLatinUp.index(text[ind])]
        elif flag == 2 and text[ind] in ArraySpecialUp:
            binary_string += ArrayBytesCod[ArraySpecialUp.index(text[ind])]
        elif text[ind] == ' ':
            binary_string += '00100'
        elif text[ind] == '\n':
            binary_string += '00010'
        else:
            flag = 0 if text[ind] in ArrayRussianUp else 1 if text[ind] in ArrayLatinUp else 2
            binary_string += '00000' if flag==0 else '11111' if flag==1 else '11011'
            continue
        ind += 1
    return binary_string


def encode_on_ascii(text):
    binary_text = ''
    for i in text:
        if i in russian_encoding_ascii:
            binary_text += bin(russian_encoding_ascii[i])[2:]
        else:
            bin_sim = bin(ord(i))[2:]
            binary_text += '0' * (8 - len(bin_sim)) + bin_sim
    return binary_text


def text_to_binary(method, text):
    if method == 'ASCII':
        binary_text = encode_on_ascii(text)
    elif method == 'MTK-2':
        binary_text = mtk2_encode(text.upper())
    else:
        binary_text = ''.join(format(byte, '08b') for byte in text.encode(method))
    return binary_text


namespace = {
    'ns0': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}
names = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def inf_o_nab_simv(run):
    r_element = run._element
    color_elem = r_element.find('.//ns0:color', namespaces=namespace)
    color = color_elem.get('{'+names+'}val') if color_elem is not None else '000000'
    shd_elem = r_element.find('.//ns0:shd', namespaces=namespace)
    background_color = shd_elem.get('{'+names+'}fill') if shd_elem is not None else 'FFFFFF'
    sz_elem = r_element.find('.//ns0:sz', namespaces=namespace)
    font_size = int(sz_elem.get('{'+names+'}val')) / 2 if sz_elem is not None else 'не установлен'
    w_elem = r_element.find('.//ns0:w', namespaces=namespace)
    scale = int(w_elem.get('{'+names+'}val')) if w_elem is not None else 100
    spacing_elem = r_element.find('.//ns0:spacing', namespaces=namespace)
    spacing = int(spacing_elem.get('{'+names+'}val')) / 2 if spacing_elem is not None else 0
    rFonts_elem = r_element.find('.//ns0:rFonts', namespaces=namespace)
    font_name = rFonts_elem.get('{' + names + '}ascii') if rFonts_elem is not None else 'не установлен'
    return [hex_decimal(color), hex_decimal(background_color), font_size, scale, spacing, font_name]

def hex_decimal(hex_string):
    if len(hex_string) == 6 and all(c in '0123456789ABCDEFabcdef' for c in hex_string):
        return ' '.join(str(int(hex_string[i:i + 2], 16)) for i in range(0, 6, 2))
    else:
        decimal_numbers = hex_string.split()
        if len(decimal_numbers) == 3 and all(num.isdigit() for num in decimal_numbers):
            return ''.join(f"{int(num):02X}" for num in decimal_numbers)
        else:
            raise ValueError("Входная строка должна быть либо 6-символьной шестнадцатеричной, либо состоять из 3 десятичных чисел.")

def apply_style(new_run, style):
    r_element = new_run._element
    if r_element.rPr is None:
        rPr = OxmlElement('w:rPr')
        r_element.append(rPr)
    else:
        rPr = r_element.rPr

    spacing_elem = OxmlElement('w:spacing')
    spacing_elem.set('{' + names + '}val', str(int(style[4]) * 2))
    rPr.append(spacing_elem)

    scale_elem = OxmlElement('w:w')
    scale_elem.set('{' + names + '}val', str(style[3]))
    rPr.append(scale_elem)

    back_elem = OxmlElement('w:shd')
    back_elem.set('{' + names + '}fill', hex_decimal(style[1]))
    rPr.append(back_elem)

    color_elem = OxmlElement('w:color')
    color_elem.set('{' + names + '}val', hex_decimal(style[0]))
    rPr.append(color_elem)


def encode_in_file(binary_string, input_file, output_file, params):
    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"Ошибка при открытии документа: {e}")
        return

    new_doc = Document()
    zagolovok_style = inf_o_nab_simv(doc.paragraphs[0].runs[0])
    text_style = next((inf_o_nab_simv(paragraph.runs[0]) for paragraph in doc.paragraphs[1:] if paragraph.runs), None)

    if text_style is None:
        print("Не удалось получить стиль текста.")
        return
    print(text_style)
    ind = 0
    new_paragraph = new_doc.add_paragraph()

    for run in doc.paragraphs[0].runs:
        for char in run.text:
            new_run = new_paragraph.add_run(char)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.strike = run.font.strike
            izm_style = zagolovok_style.copy()

            if ind < len(binary_string) and binary_string[ind] == '1':
                if params[0] == 'size':
                    izm_style[2] = (text_style[2]-params[1])+izm_style[2]
                elif params[0] == 'color':
                    izm_style[0] = params[1]
                elif params[0] == 'background_color':
                    izm_style[1] = params[1]
                elif params[0] == 'scale':
                    izm_style[3] = params[1]
                elif params[0] == 'spacing':
                    izm_style[4] = params[1]

            new_run.font.size = Pt(izm_style[2])
            new_run.font.name = izm_style[-1]
            apply_style(new_run, izm_style)
            ind += 1

    for paragraph in doc.paragraphs[1:]:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            for char in run.text:
                new_run = new_paragraph.add_run(char)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.strike = run.font.strike
                izm_style = text_style.copy()

                if ind < len(binary_string) and binary_string[ind] == '1':
                    if params[0] == 'size':
                        izm_style[2] = params[1]
                    elif params[0] == 'color':
                        izm_style[0] = params[1]
                    elif params[0] == 'background_color':
                        izm_style[1] = params[1]
                    elif params[0] == 'scale':
                        izm_style[3] = params[1]
                    elif params[0] == 'spacing':
                        izm_style[4] = params[1]

                new_run.font.size = Pt(izm_style[2])
                new_run.font.name = izm_style[-1]
                apply_style(new_run, izm_style)
                ind += 1

    try:
        new_doc.save(output_file)
    except PermissionError as e:
        print(f"Ошибка при сохранении документа: {e}")
        return

    os.startfile(output_file)


def encoding_text():
    text = input('Введите текст для кодирования: ')
    print("Выберите кодировку:")
    print("1. Бодо (МТК-2)\n2. КОИ-8R\n3. cp866\n4. Windows-1251\n5. ASCII")

    while True:
        choice1 = int(input("Введите номер желаемой кодировки (1-5): "))
        if choice1 == 1:
            encoding = 'MTK-2'
            break
        elif choice1 == 2:
            encoding = 'koi8-r'
            break
        elif choice1 == 3:
            encoding = 'cp866'
            break
        elif choice1 == 4:
            encoding = 'windows-1251'
            break
        elif choice1 == 5:
            encoding = 'ASCII'
            break

    print(
        "Выберите способ сокрытия информации:\n1. Цвет символов\n2. Цвет фона\n3. Размер шрифта\n4. Масштаб шрифта\n5. Межсимвольный интервал")
    while True:
        choice2 = int(input("Введите номер способа сокрытия информации (1-5): "))
        if choice2 == 1:
            method = 'color'
            value = input("Введите цвет текста в формате R G B (0-255 0-255 0-255): ")
            break
        elif choice2 == 2:
            method = 'background_color'
            value = input("Введите цвет фона в формате R G B (0-255 0-255 0-255): ")
            break
        elif choice2 == 3:
            method = 'size'
            value = int(input("Введите размер текста в пт: "))
            break
        elif choice2 == 4:
            method = 'scale'
            value = float(input("Введите масштаб текста (например, 99): "))
            break
        elif choice2 == 5:
            method = 'spacing'
            value = int(input("Введите межсимвольный интервал в пт (например, 1): "))
            break
    params = (method, value)
    binary_string=text_to_binary(encoding,text)
    print(binary_string)
    input_file = 'cont/2.docx'
    output_file = 'result.docx'
    encode_in_file(binary_string, input_file, output_file, params)

    #print(binary_string)

encoding_text()
